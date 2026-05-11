#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ルビふりツール - Webアプリ版（Streamlit）
"""

import io
import sys
from datetime import datetime
from pathlib import Path
from copy import deepcopy

import streamlit as st
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sudachipy import tokenizer, dictionary

# ────────────────────────────────────────────────
# ページ設定
# ────────────────────────────────────────────────

st.set_page_config(
    page_title="自動ルビふりツール",
    page_icon="✏️",
    layout="centered",
)

st.markdown("""
<style>
.rubifuri-title-img {
    filter: drop-shadow(0px 2px 4px rgba(0,0,0,0.25));
}
@media (prefers-color-scheme: dark) {
    .rubifuri-title-img {
        filter: drop-shadow(0px 0px 6px rgba(255,255,255,0.15)) brightness(1.1);
    }
}
</style>
<div style="display:flex; align-items:center; gap:16px; margin-bottom:4px;">
  <div class="rubifuri-title-img">
  <svg width="90" height="90" viewBox="140 40 400 450" xmlns="http://www.w3.org/2000/svg">
    <line x1="308" y1="168" x2="295" y2="205" stroke="#fff" stroke-width="2" stroke-dasharray="4,3"/>
    <line x1="372" y1="168" x2="385" y2="205" stroke="#fff" stroke-width="2" stroke-dasharray="4,3"/>
    <rect x="225" y="205" width="230" height="210" rx="28" fill="#FFD93D" stroke="white" stroke-width="5"/>
    <path d="M228 270 Q175 280 168 330" fill="none" stroke="#FFD93D" stroke-width="20" stroke-linecap="round"/>
    <path d="M228 270 Q175 280 168 330" fill="none" stroke="white" stroke-width="5" stroke-linecap="round"/>
    <circle cx="167" cy="337" r="16" fill="#FFD93D" stroke="white" stroke-width="5"/>
    <path d="M452 260 Q500 230 510 185" fill="none" stroke="#FFD93D" stroke-width="20" stroke-linecap="round"/>
    <path d="M452 260 Q500 230 510 185" fill="none" stroke="white" stroke-width="5" stroke-linecap="round"/>
    <circle cx="512" cy="178" r="16" fill="#FFD93D" stroke="white" stroke-width="5"/>
    <ellipse cx="300" cy="278" rx="20" ry="22" fill="#333"/>
    <ellipse cx="380" cy="278" rx="20" ry="22" fill="#333"/>
    <ellipse cx="307" cy="271" rx="7" ry="8" fill="white"/>
    <ellipse cx="387" cy="271" rx="7" ry="8" fill="white"/>
    <ellipse cx="272" cy="308" rx="20" ry="13" fill="#FF9A8B" opacity="0.55"/>
    <ellipse cx="408" cy="308" rx="20" ry="13" fill="#FF9A8B" opacity="0.55"/>
    <path d="M302 328 Q340 355 378 328" fill="none" stroke="#333" stroke-width="3.5" stroke-linecap="round"/>
    <rect x="268" y="408" width="42" height="62" rx="21" fill="#FFD93D" stroke="white" stroke-width="5"/>
    <rect x="370" y="408" width="42" height="62" rx="21" fill="#FFD93D" stroke="white" stroke-width="5"/>
    <ellipse cx="289" cy="469" rx="30" ry="16" fill="#555" stroke="white" stroke-width="4"/>
    <ellipse cx="391" cy="469" rx="30" ry="16" fill="#555" stroke="white" stroke-width="4"/>
    <circle cx="340" cy="108" r="62" fill="#FF6B9D" stroke="white" stroke-width="5"/>
    <ellipse cx="320" cy="97" rx="11" ry="13" fill="#333"/>
    <ellipse cx="360" cy="97" rx="11" ry="13" fill="#333"/>
    <ellipse cx="326" cy="91" rx="4.5" ry="5.5" fill="white"/>
    <ellipse cx="366" cy="91" rx="4.5" ry="5.5" fill="white"/>
    <ellipse cx="303" cy="112" rx="12" ry="8" fill="#FF9A8B" opacity="0.55"/>
    <ellipse cx="377" cy="112" rx="12" ry="8" fill="#FF9A8B" opacity="0.55"/>
    <path d="M320 125 Q340 140 360 125" fill="none" stroke="#333" stroke-width="2.5" stroke-linecap="round"/>
    <ellipse cx="272" cy="118" rx="20" ry="14" fill="#FF6B9D" stroke="white" stroke-width="4"/>
    <ellipse cx="410" cy="95" rx="20" ry="14" fill="#FF6B9D" stroke="white" stroke-width="4" transform="rotate(-25 410 95)"/>
  </svg>
  </div>
  <div>
    <h1 style="margin:0; font-size:2rem; font-weight:500;">自動ルビふりツール</h1>
    <p style="margin:0; color:gray; font-size:0.9rem;">Word（.docx）ファイルの漢字にルビを自動付与します</p>
  </div>
</div>
""", unsafe_allow_html=True)

# ────────────────────────────────────────────────
# SudachiPy 初期化（キャッシュ）
# ────────────────────────────────────────────────

@st.cache_resource
def load_tokenizer():
    for dict_type in ("full", "core", "small"):
        try:
            tok = dictionary.Dictionary(dict_type=dict_type).create()
            return tok
        except Exception:
            continue
    raise RuntimeError("SudachiPy辞書の読み込みに失敗しました。")

# ────────────────────────────────────────────────
# ルビ書式パラメータ
# ────────────────────────────────────────────────

DEFAULT_BASE_TEXT_SIZE = 24
DEFAULT_RUBY_FONT = "游明朝"

# ────────────────────────────────────────────────
# 横書き：ルビ書式パラメータ
# ────────────────────────────────────────────────

def get_ruby_params(sz_hpt, szCs_hpt, doc_default_hpt=24):
    if sz_hpt is not None:
        base = sz_hpt
    elif szCs_hpt is not None:
        base = szCs_hpt
    else:
        base = doc_default_hpt
    hps_base_text = base
    hps = max(8, base // 2)
    if base <= 20:
        hps_raise = 20
    elif base <= 24:
        hps_raise = 24
    else:
        hps_raise = max(24, int(base * 1.0))
    return hps, hps_raise, hps_base_text

# ────────────────────────────────────────────────
# 縦書き：ルビ書式パラメータ
# ────────────────────────────────────────────────

def get_ruby_params_tate(sz_hpt, szCs_hpt, doc_default_hpt=24, hps_raise_tate=17):
    """
    縦書き用ルビパラメータ。実ファイル解析より：
      hpsBaseText = sz（またはフォールバック）
      hps         = ceil(szCs / 2)  szCsがNoneならsz//2
      hpsRaise    = 20固定
    """
    sz   = sz_hpt   if sz_hpt   is not None else doc_default_hpt
    szCs = szCs_hpt if szCs_hpt is not None else sz  # szCsがなければszで代用

    hps_base_text = sz
    hps = max(8, -(-szCs // 2))  # ceil(szCs/2)
    hps_raise = hps_raise_tate  # 縦書き：UIから調整可能
    rt_sz   = max(7, szCs // 3)
    rt_szCs = szCs
    return hps, hps_raise, hps_base_text, rt_sz, rt_szCs


def get_run_color(rpr_elem):
    """<w:rPr> から文字色を取得する。戻り値: 16進カラーコード文字列 or None"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if rpr_elem is None:
        return None
    color_elem = rpr_elem.find(f"{{{W}}}color")
    if color_elem is not None:
        val = color_elem.get(f"{{{W}}}val")
        if val and val.upper() != "AUTO":
            return val
    return None

# ────────────────────────────────────────────────
# 漢字判定
# ────────────────────────────────────────────────

def is_kanji(char):
    cp = ord(char)
    return (
        (0x4E00 <= cp <= 0x9FFF) or
        (0x3400 <= cp <= 0x4DBF) or
        (0x20000 <= cp <= 0x2A6DF) or
        (0xF900 <= cp <= 0xFAFF)
    )

def contains_kanji(text):
    return any(is_kanji(c) for c in text)

# ────────────────────────────────────────────────
# 読み取得・送り仮名分離
# ────────────────────────────────────────────────

def kata_to_hira(text):
    result = ""
    for ch in text:
        cp = ord(ch)
        if 0x30A1 <= cp <= 0x30F6:
            result += chr(cp - 0x60)
        else:
            result += ch
    return result

def is_non_kanji_char(ch):
    """漢字以外の文字か判定（ひらがな・カタカナ・英数字・記号など）"""
    return not is_kanji(ch)

def split_surface_reading(surface, reading):
    """
    表層形と読みを「漢字連続部分」と「非漢字部分」に分離し、
    それぞれに対応する読みを割り当てる。

    例:
      "炭酸水素ナトリウム" / "たんさんすいそナトリウム"
        → [("炭酸水素", "たんさんすいそ"), ("ナトリウム", None)]

      "書く" / "かく"
        → [("書", "か"), ("く", None)]

      "二酸化炭素" / "にさんかたんそ"
        → [("二酸化炭素", "にさんかたんそ")]
    """
    # surface を「漢字連続」と「非漢字連続」に分割
    segments_surf = []
    i = 0
    while i < len(surface):
        if is_kanji(surface[i]):
            j = i
            while j < len(surface) and is_kanji(surface[j]):
                j += 1
            segments_surf.append((surface[i:j], True))
            i = j
        else:
            j = i
            while j < len(surface) and not is_kanji(surface[j]):
                j += 1
            segments_surf.append((surface[i:j], False))
            i = j

    # 非漢字部分が読みに含まれているか確認して読みを分配する
    result = []
    remaining_reading = reading

    for seg_text, is_k in segments_surf:
        if not is_k:
            # 非漢字部分（カタカナ・ひらがな等）
            # 読みの先頭または末尾から対応部分を除去
            seg_hira = kata_to_hira(seg_text)
            if remaining_reading.startswith(seg_hira):
                remaining_reading = remaining_reading[len(seg_hira):]
            elif remaining_reading.endswith(seg_hira):
                remaining_reading = remaining_reading[: len(remaining_reading) - len(seg_hira)]
            result.append((seg_text, None))
        else:
            # 漢字部分：残りの読みのうち次の非漢字セグメントの読みを除いた部分を割り当て
            # 次の非漢字セグメントを先読みして読みの末尾から除去
            next_non_kanji = ""
            idx = segments_surf.index((seg_text, True))
            for next_seg, next_is_k in segments_surf[idx + 1:]:
                if not next_is_k:
                    next_non_kanji = kata_to_hira(next_seg)
                    break

            if next_non_kanji and remaining_reading.endswith(next_non_kanji):
                kanji_reading = remaining_reading[: len(remaining_reading) - len(next_non_kanji)]
                remaining_reading = next_non_kanji
            else:
                kanji_reading = remaining_reading
                remaining_reading = ""

            if kanji_reading and kanji_reading != seg_text:
                result.append((seg_text, kanji_reading))
            else:
                result.append((seg_text, None))

    return result


def get_ruby_segments(text, tok):
    mode = tokenizer.Tokenizer.SplitMode.C
    morphemes = tok.tokenize(text, mode)
    segments = []

    for m in morphemes:
        surface = m.surface()
        if not contains_kanji(surface):
            segments.append((surface, None))
            continue
        reading = m.reading_form()
        if not reading:
            segments.append((surface, None))
            continue
        hira = kata_to_hira(reading)

        # 漢字と非漢字が混在する場合（例：炭酸水素ナトリウム）を分離
        sub_segs = split_surface_reading(surface, hira)
        segments.extend(sub_segs)

    return segments

# ────────────────────────────────────────────────
# フォントサイズ・情報取得
# ────────────────────────────────────────────────

def get_run_sz_szcs(run_elem):
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def read_pair(rpr_elem):
        if rpr_elem is None:
            return None, None
        sz = szCs = None
        sz_e = rpr_elem.find(f"{{{W}}}sz")
        if sz_e is not None:
            v = sz_e.get(f"{{{W}}}val")
            if v:
                sz = int(v)
        szCs_e = rpr_elem.find(f"{{{W}}}szCs")
        if szCs_e is not None:
            v = szCs_e.get(f"{{{W}}}val")
            if v:
                szCs = int(v)
        return sz, szCs
    rpr = run_elem.find(f"{{{W}}}rPr")
    sz, szCs = read_pair(rpr)
    if sz is not None or szCs is not None:
        return sz, szCs
    parent = run_elem.getparent()
    if parent is not None:
        ppr = parent.find(f"{{{W}}}pPr")
        if ppr is not None:
            sz, szCs = read_pair(ppr.find(f"{{{W}}}rPr"))
            if sz is not None or szCs is not None:
                return sz, szCs
    return None, None

def get_theme_fonts(doc):
    """
    theme1.xml から日本語フォント（majorFont/minorFont の Jpan）を取得する。
    戻り値: {'major': str or None, 'minor': str or None}
    """
    A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    result = {'major': None, 'minor': None}
    try:
        theme_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme'
        )
        root = etree.fromstring(theme_part.blob)
        for kind in ('major', 'minor'):
            tag = 'majorFont' if kind == 'major' else 'minorFont'
            font_elem = root.find(f".//{{{A}}}{tag}")
            if font_elem is not None:
                for f in font_elem.findall(f"{{{A}}}font"):
                    if f.get('script') == 'Jpan':
                        result[kind] = f.get('typeface')
                        break
                # Jpanがなければlatinを使用
                if result[kind] is None:
                    latin = font_elem.find(f"{{{A}}}latin")
                    if latin is not None:
                        result[kind] = latin.get('typeface')
    except Exception:
        pass
    return result


def resolve_font_from_style(style_id, styles_elem, theme_fonts=None, visited=None):
    """
    スタイル定義を再帰的に辿ってフォント情報を解決する。
    basedOn による継承チェーンも辿る。
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if not style_id or styles_elem is None:
        return None
    if visited is None:
        visited = set()
    if style_id in visited:
        return None
    visited.add(style_id)

    for style in styles_elem.findall(f"{{{W}}}style"):
        if style.get(f"{{{W}}}styleId") != style_id:
            continue
        rpr = style.find(f".//{{{W}}}rPr")
        if rpr is not None:
            rfonts = rpr.find(f"{{{W}}}rFonts")
            if rfonts is not None:
                ea = rfonts.get(f"{{{W}}}eastAsia")
                ea_theme = rfonts.get(f"{{{W}}}eastAsiaTheme")
                if ea:
                    return {'eastAsia': ea,
                            'ascii': rfonts.get(f"{{{W}}}ascii"),
                            'hAnsi': rfonts.get(f"{{{W}}}hAnsi")}
                elif ea_theme and theme_fonts:
                    resolved = None
                    if 'minor' in ea_theme.lower():
                        resolved = theme_fonts.get('minor')
                    elif 'major' in ea_theme.lower():
                        resolved = theme_fonts.get('major')
                    if resolved:
                        return {'eastAsia': resolved, 'ascii': resolved, 'hAnsi': resolved}
        # basedOn を再帰的に辿る
        based_on = style.find(f"{{{W}}}basedOn")
        if based_on is not None:
            parent_id = based_on.get(f"{{{W}}}val")
            result = resolve_font_from_style(parent_id, styles_elem, theme_fonts, visited)
            if result:
                return result
    return None


def get_style_font_info_by_id(style_id, styles_elem, theme_fonts=None):
    """スタイル定義からフォント情報を取得する（後方互換用）"""
    return resolve_font_from_style(style_id, styles_elem, theme_fonts)


def get_run_font_info(rpr_elem, theme_fonts=None, styles_elem=None):
    """
    <w:rPr> からフォント情報を返す。
    解決優先順位：
      1. run直接のrFonts
      2. rStyleのスタイル定義（basedOn継承含む）
      3. テーマフォント
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    info = {'eastAsia': None, 'ascii': None, 'hAnsi': None, 'kern': None, 'spacing': None}
    if rpr_elem is None:
        return info

    # rStyleを先に取得
    r_style_id = None
    rs = rpr_elem.find(f"{{{W}}}rStyle")
    if rs is not None:
        r_style_id = rs.get(f"{{{W}}}val")

    # 1. run直接のrFonts
    rfonts = rpr_elem.find(f"{{{W}}}rFonts")
    if rfonts is not None:
        ea = rfonts.get(f"{{{W}}}eastAsia")
        ea_theme = rfonts.get(f"{{{W}}}eastAsiaTheme")
        if ea:
            info['eastAsia'] = ea
        elif ea_theme and theme_fonts:
            if 'minor' in ea_theme.lower():
                info['eastAsia'] = theme_fonts.get('minor')
            elif 'major' in ea_theme.lower():
                info['eastAsia'] = theme_fonts.get('major')
        info['ascii'] = rfonts.get(f"{{{W}}}ascii")
        info['hAnsi'] = rfonts.get(f"{{{W}}}hAnsi")

    # 2. フォント未解決ならrStyleの継承チェーンを辿る
    if info['eastAsia'] is None and r_style_id and styles_elem is not None:
        style_font = resolve_font_from_style(r_style_id, styles_elem, theme_fonts)
        if style_font:
            info['eastAsia'] = style_font['eastAsia']
            if not info['ascii']:
                info['ascii'] = style_font.get('ascii')
            if not info['hAnsi']:
                info['hAnsi'] = style_font.get('hAnsi')

    kern = rpr_elem.find(f"{{{W}}}kern")
    if kern is not None:
        info['kern'] = kern.get(f"{{{W}}}val")
    spacing = rpr_elem.find(f"{{{W}}}spacing")
    if spacing is not None:
        info['spacing'] = spacing.get(f"{{{W}}}val")
    return info

# ────────────────────────────────────────────────
# ルビ要素生成
# ────────────────────────────────────────────────

def make_ruby_element(base_text, ruby_text, sz_hpt, szCs_hpt, doc_default_hpt, rpr_elem=None, theme_fonts=None, color_mode="black", styles_elem=None):
    hps, hps_raise, hps_base_text = get_ruby_params(sz_hpt, szCs_hpt, doc_default_hpt)
    font_info = get_run_font_info(rpr_elem, theme_fonts, styles_elem)

    ea_font = font_info['eastAsia'] or DEFAULT_RUBY_FONT
    ruby_font  = ea_font
    ruby_ascii = font_info['ascii'] or ea_font
    ruby_hansi = font_info['hAnsi'] or ea_font

    # ルビ色の決定
    ruby_color = None
    if color_mode == "match":
        ruby_color = get_run_color(rpr_elem)  # 本文と同じ色（Noneなら黒）

    ruby = OxmlElement("w:ruby")
    ruby_pr = OxmlElement("w:rubyPr")
    ruby_align = OxmlElement("w:rubyAlign")
    ruby_align.set(qn("w:val"), "distributeSpace")
    ruby_pr.append(ruby_align)
    hps_elem = OxmlElement("w:hps")
    hps_elem.set(qn("w:val"), str(hps))
    ruby_pr.append(hps_elem)
    hps_raise_elem = OxmlElement("w:hpsRaise")
    hps_raise_elem.set(qn("w:val"), str(hps_raise))
    ruby_pr.append(hps_raise_elem)
    hps_base_elem = OxmlElement("w:hpsBaseText")
    hps_base_elem.set(qn("w:val"), str(hps_base_text))
    ruby_pr.append(hps_base_elem)
    lid = OxmlElement("w:lid")
    lid.set(qn("w:val"), "ja-JP")
    ruby_pr.append(lid)
    ruby.append(ruby_pr)

    rt = OxmlElement("w:rt")
    rt_run = OxmlElement("w:r")
    rt_rpr = OxmlElement("w:rPr")
    rt_rfonts = OxmlElement("w:rFonts")
    rt_rfonts.set(qn("w:ascii"),    ruby_ascii)
    rt_rfonts.set(qn("w:hAnsi"),    ruby_hansi)
    rt_rfonts.set(qn("w:eastAsia"), ruby_font)
    rt_rpr.append(rt_rfonts)
    if font_info['kern'] is not None:
        rt_kern = OxmlElement("w:kern")
        rt_kern.set(qn("w:val"), font_info['kern'])
        rt_rpr.append(rt_kern)
    if font_info['spacing'] is not None:
        rt_spacing = OxmlElement("w:spacing")
        rt_spacing.set(qn("w:val"), font_info['spacing'])
        rt_rpr.append(rt_spacing)
    # 色指定
    if ruby_color:
        rt_color = OxmlElement("w:color")
        rt_color.set(qn("w:val"), ruby_color)
        rt_rpr.append(rt_color)
    rt_sz = OxmlElement("w:sz")
    rt_sz.set(qn("w:val"), str(hps))
    rt_rpr.append(rt_sz)
    rt_sz_cs = OxmlElement("w:szCs")
    rt_sz_cs.set(qn("w:val"), str(hps))
    rt_rpr.append(rt_sz_cs)
    rt_run.append(rt_rpr)
    rt_t = OxmlElement("w:t")
    rt_t.text = ruby_text
    if ruby_text.startswith(" ") or ruby_text.endswith(" "):
        rt_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    rt_run.append(rt_t)
    rt.append(rt_run)
    ruby.append(rt)

    ruby_base_elem = OxmlElement("w:rubyBase")
    base_run = OxmlElement("w:r")
    if rpr_elem is not None:
        try:
            base_run.append(deepcopy(rpr_elem))
        except Exception:
            pass
    base_t = OxmlElement("w:t")
    base_t.text = base_text
    if base_text.startswith(" ") or base_text.endswith(" "):
        base_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    base_run.append(base_t)
    ruby_base_elem.append(base_run)
    ruby.append(ruby_base_elem)
    return ruby


def make_ruby_element_tate(base_text, ruby_text, sz_hpt, szCs_hpt, doc_default_hpt, rpr_elem=None, theme_fonts=None, color_mode="black", hps_raise_tate=17):
    """
    縦書き用ルビ要素生成。
    実ファイル解析より：rt に sz/szCs を指定しない（Wordがhpsから自動計算）
    hpsBaseText = sz、hps = ceil(szCs/2)、hpsRaise = UIで調整可能
    """
    hps, hps_raise, hps_base_text, rt_sz, rt_szCs = get_ruby_params_tate(sz_hpt, szCs_hpt, doc_default_hpt, hps_raise_tate)
    font_info = get_run_font_info(rpr_elem, theme_fonts)

    ea_font    = font_info['eastAsia'] or DEFAULT_RUBY_FONT
    ruby_font  = ea_font
    ruby_ascii = font_info['ascii'] or ea_font
    ruby_hansi = font_info['hAnsi'] or ea_font

    ruby_color = None
    if color_mode == "match":
        ruby_color = get_run_color(rpr_elem)

    ruby = OxmlElement("w:ruby")
    ruby_pr = OxmlElement("w:rubyPr")
    ruby_align = OxmlElement("w:rubyAlign")
    ruby_align.set(qn("w:val"), "distributeSpace")
    ruby_pr.append(ruby_align)
    hps_elem = OxmlElement("w:hps")
    hps_elem.set(qn("w:val"), str(hps))
    ruby_pr.append(hps_elem)
    hps_raise_elem = OxmlElement("w:hpsRaise")
    hps_raise_elem.set(qn("w:val"), str(hps_raise))
    ruby_pr.append(hps_raise_elem)
    hps_base_elem = OxmlElement("w:hpsBaseText")
    hps_base_elem.set(qn("w:val"), str(hps_base_text))
    ruby_pr.append(hps_base_elem)
    lid = OxmlElement("w:lid")
    lid.set(qn("w:val"), "ja-JP")
    ruby_pr.append(lid)
    ruby.append(ruby_pr)

    rt = OxmlElement("w:rt")
    rt_run = OxmlElement("w:r")
    rt_rpr = OxmlElement("w:rPr")
    rt_rfonts = OxmlElement("w:rFonts")
    rt_rfonts.set(qn("w:ascii"),    ruby_ascii)
    rt_rfonts.set(qn("w:hAnsi"),    ruby_hansi)
    rt_rfonts.set(qn("w:eastAsia"), ruby_font)
    rt_rpr.append(rt_rfonts)
    if ruby_color:
        rt_color = OxmlElement("w:color")
        rt_color.set(qn("w:val"), ruby_color)
        rt_rpr.append(rt_color)
    # 縦書きは sz/szCs を指定しない（Wordがhpsから自動計算）
    rt_run.append(rt_rpr)
    rt_t = OxmlElement("w:t")
    rt_t.text = ruby_text
    if ruby_text.startswith(" ") or ruby_text.endswith(" "):
        rt_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    rt_run.append(rt_t)
    rt.append(rt_run)
    ruby.append(rt)

    ruby_base_elem = OxmlElement("w:rubyBase")
    base_run = OxmlElement("w:r")
    if rpr_elem is not None:
        try:
            base_run.append(deepcopy(rpr_elem))
        except Exception:
            pass
    base_t = OxmlElement("w:t")
    base_t.text = base_text
    if base_text.startswith(" ") or base_text.endswith(" "):
        base_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    base_run.append(base_t)
    ruby_base_elem.append(base_run)
    ruby.append(ruby_base_elem)
    return ruby


def detect_text_direction(file_bytes):
    """
    docx ファイルの書き方向を自動検出する。
    sectPr の textDirection が tbRl なら縦書き。
    戻り値: 'tate' or 'yoko'
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        doc = Document(io.BytesIO(file_bytes))
        body = doc.element.body
        sect = body.find(f"{{{W}}}sectPr")
        if sect is not None:
            td = sect.find(f"{{{W}}}textDirection")
            if td is not None:
                val = td.get(f"{{{W}}}val", "")
                if val in ("tbRl", "tbLr", "btLr"):
                    return "tate"
    except Exception:
        pass
    return "yoko"

# ────────────────────────────────────────────────
# 段落・ファイル処理
# ────────────────────────────────────────────────

def process_run(run, tok, doc_default_hpt=DEFAULT_BASE_TEXT_SIZE, theme_fonts=None, color_mode="black", tate=False, hps_raise_tate=17, styles_elem=None):
    text = run.text
    if not text or not contains_kanji(text):
        return None
    sz_hpt, szCs_hpt = get_run_sz_szcs(run._r)
    rpr_elem = run._r.find(qn("w:rPr"))
    segments = get_ruby_segments(text, tok)
    new_elements = []
    for seg_text, reading in segments:
        if reading is not None:
            if tate:
                ruby_elem = make_ruby_element_tate(seg_text, reading, sz_hpt, szCs_hpt, doc_default_hpt, rpr_elem, theme_fonts, color_mode, hps_raise_tate)
            else:
                ruby_elem = make_ruby_element(seg_text, reading, sz_hpt, szCs_hpt, doc_default_hpt, rpr_elem, theme_fonts, color_mode, styles_elem)
            new_elements.append(ruby_elem)
        else:
            plain_run = deepcopy(run._r)
            for t in plain_run.findall(qn("w:t")):
                t.text = seg_text
                if seg_text.startswith(" ") or seg_text.endswith(" "):
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_elements.append(plain_run)
    return new_elements

def apply_ruby_to_paragraph(para, tok, doc_default_hpt=DEFAULT_BASE_TEXT_SIZE, theme_fonts=None, color_mode="black", tate=False, hps_raise_tate=17, styles_elem=None):
    for run in para.runs:
        new_elems = process_run(run, tok, doc_default_hpt, theme_fonts, color_mode, tate, hps_raise_tate, styles_elem)
        if new_elems is None:
            continue
        r_elem = run._r
        parent = r_elem.getparent()
        if parent is None:
            continue
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        for i, elem in enumerate(new_elems):
            parent.insert(idx + i, elem)

def apply_ruby_to_textboxes(doc, tok, doc_default_hpt=DEFAULT_BASE_TEXT_SIZE, theme_fonts=None, color_mode="black", tate=False, hps_raise_tate=17, styles_elem=None):
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    from docx.text.paragraph import Paragraph as DocxParagraph
    body = doc.element.body
    txbx_contents = body.findall(f".//{{{W}}}txbxContent")
    txbx_contents += body.findall(f".//{{{MC}}}Fallback//{{{W}}}txbxContent")
    for txbx in txbx_contents:
        for p_elem in txbx.findall(f"{{{W}}}p"):
            para = DocxParagraph(p_elem, doc)
            apply_ruby_to_paragraph(para, tok, doc_default_hpt, theme_fonts, color_mode, tate, hps_raise_tate, styles_elem)

def get_doc_default_font_size(doc):
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        root = doc.part.styles._element
        sz = root.find(f".//{{{W}}}rPrDefault/{{{W}}}rPr/{{{W}}}sz")
        if sz is not None:
            val = sz.get(f"{{{W}}}val")
            if val:
                return int(val)
    except Exception:
        pass
    return DEFAULT_BASE_TEXT_SIZE

def process_docx(file_bytes, filename, tok, color_mode="black", tate=False, hps_raise_tate=17):
    doc = Document(io.BytesIO(file_bytes))
    doc_default_hpt = get_doc_default_font_size(doc)
    theme_fonts = get_theme_fonts(doc)
    styles_elem = doc.part.styles._element if doc.part.styles else None
    for para in doc.paragraphs:
        apply_ruby_to_paragraph(para, tok, doc_default_hpt, theme_fonts, color_mode, tate, hps_raise_tate, styles_elem)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    apply_ruby_to_paragraph(para, tok, doc_default_hpt, theme_fonts, color_mode, tate, hps_raise_tate, styles_elem)
    apply_ruby_to_textboxes(doc, tok, doc_default_hpt, theme_fonts, color_mode, tate, hps_raise_tate, styles_elem)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ────────────────────────────────────────────────
# 片山モード専用処理
# ────────────────────────────────────────────────

# スタイルIDごとのhpsRaiseテーブル
KATAYAMA_STYLE_HPSRAISE = {
    None:  20,   # スタイルなし → \s\up 10
    "a7":  20,   # 大問 リード文 → \s\up 10
    "a":   24,   # 設問⑴⑵⑶ → \s\up 12
    "a0":  24,   # 設問①②③ → \s\up 12
}
KATAYAMA_DEFAULT_HPSRAISE = 20

# ルビをスキップするスタイル名キーワード（出典・注など）
KATAYAMA_SKIP_STYLE_KEYWORDS = ["出典", "source", "citation", "注", "footer", "フッター"]
KATAYAMA_SKIP_STYLE_IDS = {"af0", "af1", "a5"}  # footer系・Subtle Reference（出典）


def get_para_style_id(para):
    """段落のスタイルIDを返す。なければNone。"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ppr = para._p.find(f"{{{W}}}pPr")
    if ppr is not None:
        ps = ppr.find(f"{{{W}}}pStyle")
        if ps is not None:
            return ps.get(f"{{{W}}}val")
    return None


def should_skip_ruby_katayama(para, doc):
    """片山モードでルビをスキップすべき段落かどうか判定"""
    style_id = get_para_style_id(para)

    # IDによるスキップ
    if style_id in KATAYAMA_SKIP_STYLE_IDS:
        return True

    # スタイル名によるスキップ（出典・注など）
    if style_id and doc.part.styles:
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for style in doc.part.styles._element.findall(f"{{{W}}}style"):
            sid = style.get(f"{{{W}}}styleId")
            if sid == style_id:
                name_elem = style.find(f"{{{W}}}name")
                if name_elem is not None:
                    name = name_elem.get(f"{{{W}}}val", "")
                    if any(kw in name for kw in KATAYAMA_SKIP_STYLE_KEYWORDS):
                        return True
    return False


def find_page1_end_index(doc):
    """
    lastRenderedPageBreak を持つ最初の段落のインデックスを返す。
    見つからなければ0（スキップなし）。
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for i, para in enumerate(doc.paragraphs):
        breaks = para._p.findall(f".//{{{W}}}lastRenderedPageBreak")
        br_page = para._p.findall(f".//{{{W}}}br[@{{{W}}}type='page']")
        if breaks or br_page:
            return i
    return 0


def get_run_style_id(run):
    """runの文字スタイルIDを返す"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rpr = run._r.find(f"{{{W}}}rPr")
    if rpr is not None:
        rs = rpr.find(f"{{{W}}}rStyle")
        if rs is not None:
            return rs.get(f"{{{W}}}val")
    return None


def apply_ruby_to_paragraph_katayama(para, tok, doc_default_hpt, theme_fonts, color_mode, styles_elem=None):
    """片山モード：スタイルIDに応じたhpsRaiseを適用、runスタイルのスキップも対応"""
    style_id = get_para_style_id(para)
    hps_raise = KATAYAMA_STYLE_HPSRAISE.get(style_id, KATAYAMA_DEFAULT_HPSRAISE)

    for run in para.runs:
        # runの文字スタイルがスキップ対象なら除外（例：a5=出典）
        run_style_id = get_run_style_id(run)
        if run_style_id in KATAYAMA_SKIP_STYLE_IDS:
            continue

        text = run.text
        if not text or not contains_kanji(text):
            continue
        sz_hpt, szCs_hpt = get_run_sz_szcs(run._r)
        rpr_elem = run._r.find(qn("w:rPr"))

        if sz_hpt is not None:
            base = sz_hpt
        elif szCs_hpt is not None:
            base = szCs_hpt
        else:
            base = doc_default_hpt
        hps = max(8, base // 2)
        hps_base_text = base

        # styles_elemを渡してフォント継承を解決
        font_info = get_run_font_info(rpr_elem, theme_fonts, styles_elem)
        ea_font = font_info['eastAsia'] or DEFAULT_RUBY_FONT
        ruby_color = get_run_color(rpr_elem) if color_mode == "match" else None

        segments = get_ruby_segments(text, tok)
        new_elements = []
        for seg_text, reading in segments:
            if reading is not None:
                ruby = OxmlElement("w:ruby")
                ruby_pr = OxmlElement("w:rubyPr")
                for tag, val in [("w:rubyAlign", "distributeSpace"), ("w:hps", str(hps)),
                                  ("w:hpsRaise", str(hps_raise)), ("w:hpsBaseText", str(hps_base_text)),
                                  ("w:lid", "ja-JP")]:
                    e = OxmlElement(tag)
                    e.set(qn("w:val"), val)
                    ruby_pr.append(e)
                ruby.append(ruby_pr)

                rt = OxmlElement("w:rt")
                rt_run = OxmlElement("w:r")
                rt_rpr = OxmlElement("w:rPr")
                rt_rfonts = OxmlElement("w:rFonts")
                rt_rfonts.set(qn("w:ascii"), ea_font)
                rt_rfonts.set(qn("w:hAnsi"), ea_font)
                rt_rfonts.set(qn("w:eastAsia"), ea_font)
                rt_rpr.append(rt_rfonts)
                if ruby_color:
                    rc = OxmlElement("w:color")
                    rc.set(qn("w:val"), ruby_color)
                    rt_rpr.append(rc)
                for sz_tag in ("w:sz", "w:szCs"):
                    se = OxmlElement(sz_tag)
                    se.set(qn("w:val"), str(hps))
                    rt_rpr.append(se)
                rt_run.append(rt_rpr)
                rt_t = OxmlElement("w:t")
                rt_t.text = reading
                rt_run.append(rt_t)
                rt.append(rt_run)
                ruby.append(rt)

                ruby_base = OxmlElement("w:rubyBase")
                base_run = OxmlElement("w:r")
                if rpr_elem is not None:
                    try:
                        base_run.append(deepcopy(rpr_elem))
                    except Exception:
                        pass
                base_t = OxmlElement("w:t")
                base_t.text = seg_text
                base_run.append(base_t)
                ruby_base.append(base_run)
                ruby.append(ruby_base)
                new_elements.append(ruby)
            else:
                plain_run = deepcopy(run._r)
                for t in plain_run.findall(qn("w:t")):
                    t.text = seg_text
                new_elements.append(plain_run)

        r_elem = run._r
        parent = r_elem.getparent()
        if parent is None:
            continue
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        for i, elem in enumerate(new_elements):
            parent.insert(idx + i, elem)


def get_body_elements_after_page1(doc):
    """
    bodyの直接子要素を順番に走査し、最初のページ区切り以降の要素を返す。
    テーブルを含む1ページ目スキップに対応。
    戻り値: (after_elements, page_break_found)
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc.element.body
    after = []
    found = False
    for child in body:
        if found:
            after.append(child)
        else:
            # この要素またはその子孫にページ区切りがあるか
            if (child.findall(f".//{{{W}}}lastRenderedPageBreak") or
                    child.findall(f".//{{{W}}}br[@{{{W}}}type='page']")):
                found = True
                after.append(child)  # ページ区切りを含む要素は含める
    return after, found


def process_docx_katayama(file_bytes, filename, tok, color_mode="black"):
    """片山モード：1ページ目スキップ＋スタイル別hpsRaise"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable

    doc = Document(io.BytesIO(file_bytes))
    doc_default_hpt = get_doc_default_font_size(doc)
    theme_fonts = get_theme_fonts(doc)
    styles_elem = doc.part.styles._element if doc.part.styles else None

    # 1ページ目以降のbody要素を取得
    after_elements, found = get_body_elements_after_page1(doc)

    if not found:
        # ページ区切りが見つからない場合は全体を処理
        after_elements = list(doc.element.body)

    for elem in after_elements:
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            para = DocxParagraph(elem, doc)
            if should_skip_ruby_katayama(para, doc):
                continue
            apply_ruby_to_paragraph_katayama(para, tok, doc_default_hpt, theme_fonts, color_mode, styles_elem)
        elif tag == 'tbl':
            table = DocxTable(elem, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if should_skip_ruby_katayama(para, doc):
                            continue
                        apply_ruby_to_paragraph_katayama(para, tok, doc_default_hpt, theme_fonts, color_mode, styles_elem)

    apply_ruby_to_textboxes(doc, tok, doc_default_hpt, theme_fonts, color_mode)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ────────────────────────────────────────────────
# UI
# ────────────────────────────────────────────────

# 辞書を裏で読み込み（表示なし）
try:
    tok = load_tokenizer()
except Exception as e:
    st.error(f"辞書の読み込みに失敗しました: {e}")
    st.stop()

# 片山モード：セッション状態で管理
if "katayama_count" not in st.session_state:
    st.session_state.katayama_count = 0
if "katayama_mode" not in st.session_state:
    st.session_state.katayama_mode = False

katayama_mode = st.query_params.get("mode") == "katayama"

st.divider()

uploaded_file = st.file_uploader(
    "Wordファイル（.docx）をアップロードしてください",
    type=["docx"],
    help=".docx 形式のファイルのみ対応しています"
)

if uploaded_file is not None:
    st.info(f"📄 ファイル名: {uploaded_file.name}")

    # 書き方向の自動検出
    file_bytes_preview = uploaded_file.read()
    auto_direction = detect_text_direction(file_bytes_preview)
    uploaded_file.seek(0)  # 読み直しのためにリセット

    direction_default = 0 if auto_direction == "yoko" else 1
    st.caption(f"🔍 自動検出: {'横書き' if auto_direction == 'yoko' else '縦書き'}")

    # 書き方向の選択
    direction_choice = st.radio(
        "書き方向",
        options=["yoko", "tate"],
        format_func=lambda x: "↔️ 横書き" if x == "yoko" else "↕️ 縦書き",
        index=direction_default,
        horizontal=True,
    )

    # ルビ色の選択
    color_choice = st.radio(
        "ルビの文字色",
        options=["black", "match"],
        format_func=lambda x: "⬛ すべて黒色" if x == "black" else "🎨 本文の色に合わせる",
        horizontal=True,
    )

    # 縦書きのみ：ルビ距離スライダー
    hps_raise_tate = 17
    if direction_choice == "tate":
        offset = st.slider(
            "ルビの距離調整（縦書き用）",
            min_value=-5,
            max_value=5,
            value=0,
            step=1,
            help="0がデフォルト  \n＋で離れる　－で近づく"
        )
        hps_raise_tate = 17 + offset

    # 片山モード表示
    if katayama_mode:
        st.markdown(
            "<p style='color:#FF6B9D; font-size:0.8rem;'>✨ 片山モード：スタイル別設定・１ページ目スキップ</p>",
            unsafe_allow_html=True
        )

    if st.button("✨ ルビをふる", type="primary", use_container_width=True):
        with st.spinner("処理中です..."):
            try:
                file_bytes = uploaded_file.read()
                if katayama_mode:
                    result = process_docx_katayama(file_bytes, uploaded_file.name, tok, color_mode=color_choice)
                else:
                    tate = (direction_choice == "tate")
                    result = process_docx(file_bytes, uploaded_file.name, tok, color_mode=color_choice, tate=tate, hps_raise_tate=hps_raise_tate)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                stem = Path(uploaded_file.name).stem
                out_name = f"ルビ版_{timestamp}_{stem}.docx"

                st.success("✅ 処理完了！")
                st.download_button(
                    label="📥 ダウンロード",
                    data=result,
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"エラーが発生しました: {e}")
                with st.expander("🐛 不具合を報告する"):
                    st.write("以下の内容を送信すると、開発者に通知されます。")
                    report_name = st.text_input("お名前（任意）", key="report_name")
                    report_msg = st.text_area(
                        "状況・コメント（任意）",
                        value=f"エラー内容: {e}\nファイル名: {uploaded_file.name}",
                        key="report_msg"
                    )
                    if st.button("📨 報告を送信", key="report_btn"):
                        import urllib.request, urllib.parse, json as _json
                        data = urllib.parse.urlencode({
                            "name": report_name or "匿名",
                            "message": report_msg,
                            "_subject": "【ルビふりくん】不具合報告",
                        }).encode()
                        req = urllib.request.Request(
                            "https://formspree.io/f/mykolpee",
                            data=data,
                            headers={"Accept": "application/json"},
                        )
                        try:
                            urllib.request.urlopen(req)
                            st.success("✅ 報告を送信しました。ありがとうございます！")
                        except Exception:
                            st.warning("送信に失敗しました。時間をおいて再度お試しください。")

st.divider()
st.caption("ルビの読みはSudachiPy（全辞書）を使用しています。付与後に内容をご確認ください。")

st.markdown("<div style='margin-top:48px;'></div>", unsafe_allow_html=True)

st.markdown("""
<style>
div[data-testid="stExpander"] {
    border: 1px solid #ddd !important;
    opacity: 0.65;
}
div[data-testid="stExpander"] summary {
    font-size: 0.78rem !important;
    color: #888 !important;
}
</style>
""", unsafe_allow_html=True)

with st.expander("💬 ご意見・ご要望・不具合報告"):
    st.markdown("<small>お気づきの点や改善のご要望、不具合などがあればお知らせください。<br>いただいた内容は開発者が確認し、ツールの改善に役立てます。</small>", unsafe_allow_html=True)
    feedback_name = st.text_input("お名前（任意）", key="feedback_name")
    feedback_type = st.radio(
        "種類",
        options=["不具合の報告", "改善・修正のご要望", "新機能のリクエスト", "その他"],
        horizontal=True,
        key="feedback_type"
    )
    feedback_msg = st.text_area(
        "内容",
        placeholder="例：〇〇のときにルビがうまく振られませんでした。／〇〇の機能を追加してほしいです。",
        key="feedback_msg"
    )

    if st.button("📨 送信する", key="feedback_btn"):
        if not feedback_msg.strip():
            st.warning("内容を入力してください。")
        else:
            import json as _json
            payload = _json.dumps({
                "name": feedback_name or "匿名",
                "message": f"【{feedback_type}】\n\n{feedback_msg}",
                "_subject": f"【ルビふりくん】{feedback_type}",
            })
            st.components.v1.html(f"""
            <script>
            fetch("https://formspree.io/f/mykolpee", {{
                method: "POST",
                headers: {{"Content-Type": "application/json", "Accept": "application/json"}},
                body: {payload!r}
            }})
            .then(r => r.json())
            .then(data => {{
                if (data.ok) {{
                    document.getElementById("result").innerText = "✅ 送信しました。ありがとうございます！";
                }} else {{
                    document.getElementById("result").innerText = "❌ 送信に失敗しました: " + JSON.stringify(data);
                }}
            }})
            .catch(e => {{
                document.getElementById("result").innerText = "❌ 通信エラー: " + e;
            }});
            </script>
            <p id="result" style="font-family:sans-serif; color:#444;">送信中...</p>
            """, height=40)

st.markdown(
    "<p style='text-align:right; color:#bbb; font-size:0.72rem; margin-top:8px;'>Developed by かたやま</p>",
    unsafe_allow_html=True
)

# ルビふりくんを親DOMに注入（クリック3回で片山モードON/OFF）
_svg = """
<svg width="100%" viewBox="0 0 680 500" xmlns="http://www.w3.org/2000/svg">
  <line x1="308" y1="168" x2="295" y2="205" stroke="#fff" stroke-width="2" stroke-dasharray="4,3"/>
  <line x1="372" y1="168" x2="385" y2="205" stroke="#fff" stroke-width="2" stroke-dasharray="4,3"/>
  <rect x="225" y="205" width="230" height="210" rx="28" fill="#FFD93D" stroke="white" stroke-width="5"/>
  <path d="M228 270 Q175 280 168 330" fill="none" stroke="#FFD93D" stroke-width="20" stroke-linecap="round"/>
  <path d="M228 270 Q175 280 168 330" fill="none" stroke="white" stroke-width="5" stroke-linecap="round"/>
  <circle cx="167" cy="337" r="16" fill="#FFD93D" stroke="white" stroke-width="5"/>
  <path d="M452 260 Q500 230 510 185" fill="none" stroke="#FFD93D" stroke-width="20" stroke-linecap="round"/>
  <path d="M452 260 Q500 230 510 185" fill="none" stroke="white" stroke-width="5" stroke-linecap="round"/>
  <circle cx="512" cy="178" r="16" fill="#FFD93D" stroke="white" stroke-width="5"/>
  <ellipse cx="300" cy="278" rx="20" ry="22" fill="#333"/>
  <ellipse cx="380" cy="278" rx="20" ry="22" fill="#333"/>
  <ellipse cx="307" cy="271" rx="7" ry="8" fill="white"/>
  <ellipse cx="387" cy="271" rx="7" ry="8" fill="white"/>
  <ellipse cx="272" cy="308" rx="20" ry="13" fill="#FF9A8B" opacity="0.55"/>
  <ellipse cx="408" cy="308" rx="20" ry="13" fill="#FF9A8B" opacity="0.55"/>
  <path d="M302 328 Q340 355 378 328" fill="none" stroke="#333" stroke-width="3.5" stroke-linecap="round"/>
  <rect x="268" y="408" width="42" height="62" rx="21" fill="#FFD93D" stroke="white" stroke-width="5"/>
  <rect x="370" y="408" width="42" height="62" rx="21" fill="#FFD93D" stroke="white" stroke-width="5"/>
  <ellipse cx="289" cy="469" rx="30" ry="16" fill="#555" stroke="white" stroke-width="4"/>
  <ellipse cx="391" cy="469" rx="30" ry="16" fill="#555" stroke="white" stroke-width="4"/>
  <circle cx="340" cy="108" r="62" fill="#FF6B9D" stroke="white" stroke-width="5"/>
  <ellipse cx="320" cy="97" rx="11" ry="13" fill="#333"/>
  <ellipse cx="360" cy="97" rx="11" ry="13" fill="#333"/>
  <ellipse cx="326" cy="91" rx="4.5" ry="5.5" fill="white"/>
  <ellipse cx="366" cy="91" rx="4.5" ry="5.5" fill="white"/>
  <ellipse cx="303" cy="112" rx="12" ry="8" fill="#FF9A8B" opacity="0.55"/>
  <ellipse cx="377" cy="112" rx="12" ry="8" fill="#FF9A8B" opacity="0.55"/>
  <path d="M320 125 Q340 140 360 125" fill="none" stroke="#333" stroke-width="2.5" stroke-linecap="round"/>
  <ellipse cx="272" cy="118" rx="20" ry="14" fill="#FF6B9D" stroke="white" stroke-width="4"/>
  <ellipse cx="410" cy="95" rx="20" ry="14" fill="#FF6B9D" stroke="white" stroke-width="4" transform="rotate(-25 410 95)"/>
</svg>
"""

st.components.v1.html(f"""
<script>
(function() {{
  var doc = window.parent.document;
  if (doc.getElementById('rubifuri-kun-fixed')) return;

  var div = doc.createElement('div');
  div.id = 'rubifuri-kun-fixed';
  div.style.cssText = [
    'position:fixed', 'bottom:60px', 'right:16px', 'width:120px',
    'opacity:0.88', 'z-index:9999', 'cursor:pointer',
    'filter:drop-shadow(0px 2px 4px rgba(0,0,0,0.25))',
    'transition:transform 0.1s'
  ].join(';');
  div.innerHTML = `{_svg}`;
  doc.body.appendChild(div);

  var count = 0, timer = null;

  // クリックハンドラを親コンテキストで実行
  var parentFn = window.parent.Function('div', 'count_ref', `
    return function() {{
      count_ref.v++;
      div.style.transform = 'scale(1.15)';
      setTimeout(function(){{ div.style.transform = 'scale(1)'; }}, 120);
      if (count_ref.v >= 3) {{
        count_ref.v = 0;
        var url = new URL(window.location.href);
        if (url.searchParams.get('mode') === 'katayama') {{
          url.searchParams.delete('mode');
        }} else {{
          url.searchParams.set('mode', 'katayama');
        }}
        window.location.href = url.toString();
      }}
      clearTimeout(count_ref.t);
      count_ref.t = setTimeout(function(){{ count_ref.v = 0; }}, 2000);
    }};
  `);

  var count_ref = {{v: 0, t: null}};
  div.addEventListener('click', parentFn(div, count_ref));

  if (window.parent.matchMedia('(prefers-color-scheme: dark)').matches) {{
    div.style.opacity = '1.0';
    div.style.filter = 'drop-shadow(0px 0px 6px rgba(255,255,255,0.15)) brightness(1.1)';
  }}
}})();
</script>
""", height=0)
